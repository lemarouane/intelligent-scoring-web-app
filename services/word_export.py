# ============================================================
# BOA INTELLIGENT CREDIT SCORING – WORD EXPORT
# ============================================================

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_heading_color(paragraph, hex_color: str):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )


def _add_table_row(table, cells: list, bold: bool = False, bg: str = None):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(val)
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
        run.font.bold = bold
        run.font.size = Pt(9)
        if bg:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), bg)
            tc_pr.append(shd)
    return row


def generate_word(
    categorie: str,
    final_score: float,
    risk_class: str,
    decision: str,
    prob_default: str,
    rows: list,
    ia_analysis: str = "",
) -> bytes:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title
    title = doc.add_heading("BANK OF AFRICA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, "003366")

    sub = doc.add_paragraph("Système Intelligent de Scoring de Crédit Professionnel")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)
    sub.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Meta table
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = "Table Grid"
    meta_data = [
        ("Date d'analyse", datetime.now().strftime("%d/%m/%Y %H:%M")),
        ("Catégorie client", categorie),
        ("Score Final", f"{final_score:.2f} / 100"),
        ("Classe de risque", f"Classe {risk_class}"),
        ("Décision", decision),
        
    ]
    for label, value in meta_data:
        _add_table_row(meta_table, [label, value])

    doc.add_paragraph()

    # Score heading
    h2 = doc.add_heading("Détail des critères de scoring", level=2)
    _set_heading_color(h2, "003366")

    score_table = doc.add_table(rows=1, cols=4)
    score_table.style = "Table Grid"
    hdr = score_table.rows[0].cells
    for i, txt in enumerate(["Critère", "Poids (%)", "Score Partiel", "Score Pondéré"]):
        hdr[i].text = txt
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "003366")
        tc_pr.append(shd)

    for r in rows:
        _add_table_row(
            score_table,
            [r["label"], f"{r['poids']}%", f"{r['partial']}/100", str(r["weighted"])],
        )

    doc.add_paragraph()

    # AI Analysis
    if ia_analysis:
        h2b = doc.add_heading("Analyse Intelligence Artificielle", level=2)
        _set_heading_color(h2b, "003366")
        for line in ia_analysis.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            if line.startswith("**") and line.endswith("**"):
                run = p.add_run(line[2:-2])
                run.bold = True
                run.font.size = Pt(10)
            elif line.startswith("#"):
                run = p.add_run(line.lstrip("#").strip())
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
                run.font.size = Pt(10)
            else:
                run = p.add_run(line)
                run.font.size = Pt(9)

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        "Document confidentiel – Usage interne – Bank of Africa © 2026"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(7)
    footer_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()